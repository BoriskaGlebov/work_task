import os
import shutil
import tempfile
import unittest
from docx import Document as DocxDocument

from file_creator.utils.parser_word.my_parser import Parser


class TestParser(unittest.TestCase):
    def setUp(self):
        # Создаем временную директорию для тестовых файлов
        self.test_dir = tempfile.mkdtemp()

        # Создаем 2 тестовых docx-файла с разным содержимым
        self.file1 = os.path.join(self.test_dir, "test1.docx")
        self.file2 = os.path.join(self.test_dir, "test2.docx")

        self._create_docx(self.file1, [
            "Куда и кому: Тестовый адресат",
            "Уважаемый клиент, спасибо за обращение.",
            "Просто абзац с текстом."
        ])

        self._create_docx(self.file2, [
            "Еще один тестовый документ.",
            "Evaluation Warning: The document was created with Spire.Doc for Python.",
            "Особый знак"
        ], add_table=True)

        self.start_number = 100

    def _create_docx(self, path, paragraphs, add_table=False):
        doc = DocxDocument()
        # Добавляем параграфы
        for para_text in paragraphs:
            doc.add_paragraph(para_text)

        # При необходимости добавляем таблицу с нужным текстом
        if add_table:
            table = doc.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "Особый знак"

        doc.save(path)

    def tearDown(self):
        # Удаляем временную директорию после тестов
        shutil.rmtree(self.test_dir)

    def test_all_files(self):
        parser = Parser(self.test_dir, self.start_number)
        files = parser.all_files()
        self.assertCountEqual(files, ["test1.docx", "test2.docx"])

    def test_create_file_parsed(self):
        parser = Parser(self.test_dir, self.start_number)
        result = parser.create_file_parsed()

        # Проверяем, что вернулось два результата
        self.assertEqual(len(result), 2)

        # Проверяем, что строки отформатированы и содержат ожидаемые части
        self.assertIn("ТЕСТОВЫЙ АДРЕСАТ", result[0].upper())
        self.assertIn("УВАЖАЕМЫЙ КЛИЕНТ", result[0].upper())
        self.assertIn("ПРОСТО АБЗАЦ", result[0].upper())

        self.assertIn("ОСОБЫЙ ЗНАК", result[1].upper())
        self.assertIn("ЗАМЕСТИТЕЛЬ ДОЯРКИ", result[1].upper())

