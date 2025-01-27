import datetime
import os
from typing import List

from docx import Document

from work_for_ilia.utils.my_settings.disrs_for_app import ProjectSettings


class Parser:
    """
    Парсит документ .docx и извлекает текст в отформатированном виде.

    Attributes:
        directory (str): Путь к директории, содержащей файлы .docx.
        start_number (int): Начальный номер для именования выходных файлов.
    """

    def __init__(self, directory: str, start_number: int):
        """
        Инициализирует экземпляр класса Parser.

        Args:
            directory (str): Путь к директории для обработки файлов.
            start_number (int): Начальный номер для именования выходных файлов.
        """
        self.directory = directory
        self.start_number = start_number

    def all_files(self) -> List[str]:
        """
        Находит все файлы .docx в указанной директории.

        Returns:
            list[str]: Список имен файлов .docx в директории.
        """
        files = [
            file for file in os.listdir(self.directory)
            if os.path.isfile(os.path.join(self.directory, file)) and file.endswith('.docx')
        ]
        return files

    def format_text(self, text: str, max_length: int = 60) -> str:
        """
        Форматирует текст, ограничивая длину строк до max_length символов.

        Args:
            text (str): Текст для форматирования.
            max_length (int): Максимальная длина строки.

        Returns:
            str: Отформатированный текст с ограниченной длиной строк.
        """
        words = text.split()
        formatted_lines = []
        current_line = []

        for word in words:
            # Проверяем, если добавление слова превышает максимальную длину
            if len(' '.join(current_line + [word])) <= max_length:
                current_line.append(word)
            else:
                # Сохраняем текущую строку и начинаем новую
                formatted_lines.append(' '.join(current_line))
                current_line = [word]

        # Добавляем последнюю строку, если она не пустая
        if current_line:
            formatted_lines.append(' '.join(current_line))

        return '\n'.join(formatted_lines)

    def create_file_parsed(self) -> List[str]:
        """
        Создает отредактированные файлы .txt из документов .docx.

        Returns:
            list[str]: Список содержимого отредактированных файлов.
        """
        content_files = []

        for file in self.all_files():
            document = Document(os.path.join(self.directory, file))
            n_name = f"{self.start_number}_{os.path.splitext(file)[0]}.txt"
            out_txt = ''

            # Читаем верхний колонтитул
            special_header = document.sections[0].first_page_header
            common_header = document.sections[0].header

            header = special_header if len(special_header.tables) else common_header

            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if 'из:' in cell.text.lower():
                            out_txt += self.format_text(cell.text[:-1].strip().upper()) + ' '
                        elif 'г. москва' in cell.text.lower():
                            out_txt += self.format_text(
                                cell.text.strip().upper()) + f'  НР {self.start_number}   Для анального пользования\n'.upper()

            # Читаем основной текст документа и таблицы
            out_txt += "\n\n          Содержимое документа:\n\n"
            num_tables = 0

            for element in document.element.body:
                if element.tag.endswith('p'):  # Проверяем, является ли элемент абзацем
                    text: str = element.text.strip()
                    if element.text.startswith(
                            'Evaluation Warning: The document was created with Spire.Doc for Python.'):
                        continue
                    elif element.text.startswith('Куда и кому:'):
                        new_str = text.replace('Куда и кому:', '')
                        out_txt += self.format_text(new_str.upper()) + '\n'
                    elif element.text.startswith('Уважаемый'):
                        out_txt += '      ' + self.format_text(text.upper()) + '\n'
                    elif text:  # Если абзац не пустой
                        out_txt += self.format_text(text.upper()) + '\n'
                    else:  # Печатаем пустую строку для пустого абзаца
                        out_txt += '\n'
                elif element.tag.endswith('tbl') and num_tables < len(document.tables):
                    table = document.tables[num_tables]  # Получаем таблицу по индексу
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.startswith('Особый знак'):
                                out_txt += (f"Особый знак НР {str(self.start_number)}/П Заместитель доярки\n"
                                            f"{datetime.datetime.now().strftime('%d.%m.%Y')}   колхозник   А.М. Поликарп \n")
                                break

                    num_tables += 1

            self.start_number += 1
            content_files.append(out_txt)

        return content_files


def replace_unsupported_characters(text, replacement='?'):
    """
    Заменяет неподдерживаемые символы на указанный символ замены.

    Args:
        text (str): Исходный текст.
        replacement (str): Символ, на который будут заменены неподдерживаемые символы.

    Returns:
        str: Текст с замененными неподдерживаемыми символами.
    """
    # Создаем новый текст с заменами
    filtered_text = ''.join(char if can_encode(char) else replacement for char in text)

    return filtered_text


def can_encode(char):
    """
    Проверяет, может ли символ быть закодирован в cp866.

    Args:
        char (str): Символ для проверки.

    Returns:
        bool: True, если символ может быть закодирован в cp866, иначе False.
    """
    try:
        char.encode('cp866')
        return True
    except UnicodeEncodeError:
        return False


if __name__ == '__main__':
    start_numm = 123
    s = Parser(ProjectSettings.tlg_dir, start_numm)
    print(s)
    print(s.all_files())
    print(s.create_file_parsed()[0])
    # print(os.listdir(ProjectSettings.tlg_dir))
    # os.remove(*os.listdir(ProjectSettings.tlg_dir))
