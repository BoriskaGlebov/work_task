import os
import re
from pprint import pprint

from django.db.models import Q
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Inches

# Укажите путь к настройкам вашего проекта
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "lazy_ilya.settings")
import django

# Настройка Django
django.setup()

from typing import Dict, List, Optional

from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer
from docx import Document
from work_for_ilia.models import SomeDataFromSomeTables, SomeTables
from work_for_ilia.utils.my_settings.settings_for_app import ProjectSettings, logger


# from work_for_ilia.utils.my_settings.settings_for_app import ProjectSettings


class GlobusParser:
    """
    Класс обрабатывает содержимое файла с Городами
    """

    def __init__(self):
        self.model_inf = {
            "table_id": "",
            "dock_num": "",
            "location": "",
            "name_organ": "",
            "pseudonim": "",
            "letters": "",
            "writing": "",
            "ip_address": "",
            "some_number": "",
            "work_timme": "",
        }

    @classmethod
    @logger.catch(message="Непредвиденное исключение")
    def process_file(cls, file_path: str) -> None:
        """
        Обрабатывает загруженный файл с данными о городах.
        Обновляет названия таблиц на основе номера раздела.

        Args:
            file_path (str): Путь к загруженному файлу.
        """
        doc = Document(os.path.join(ProjectSettings.tlg_dir, file_path))
        tables = doc.tables
        paragraphs = doc.paragraphs
        processed_tables = []  # Список обработанных разделов
        tables_to_add = []
        tables_to_update = []

        processed_cities = []
        cities_to_add = []
        cities_to_update = []

        for paragraph in paragraphs[1:]:
            pattern = r"^Раздел (\d+)\s*(.*)"  # Группировка номера и названия
            text = paragraph.text.strip()

            match = re.match(pattern, text)

            if match:
                section_number = match.group(1)  # Номер раздела
                section_name = text  # Полное название раздела
                # Поиск таблицы по номеру раздела (в table_name)
                table: Optional[SomeTables] = SomeTables.objects.filter(
                    Q(table_name__startswith=f"Раздел {section_number}")
                    & Q(table_name__regex=rf"^Раздел {section_number}(?!\d)")
                ).first()  # Ищем по началу строки
                if table:
                    # Обновление существующей таблицы
                    if table.table_name != section_name:
                        table.table_name = section_name
                        tables_to_update.append(table)
                        processed_tables.append(table)  # Запоминаем раздел
                    elif table and table.table_name == section_name:
                        processed_tables.append(table)  # Запоминаем раздел
                else:
                    # Создание новой таблицы
                    table = SomeTables(table_name=section_name)
                    tables_to_add.append(table)
                    processed_tables.append(table)  # Запоминаем ID
        try:
            if tables_to_add:
                res = SomeTables.objects.bulk_create(tables_to_add)
                for el in res:
                    logger.info(
                        f"Создана новая таблица '{el.table_name}' (ID: {el.id})"
                    )
            if tables_to_update:
                res2 = SomeTables.objects.bulk_update(
                    tables_to_update,
                    [
                        "table_name",
                    ],
                )
                for el in tables_to_update:
                    logger.info(
                        f"Обновлено название таблицы '{el.table_name}' (ID: {el.id})"
                    )
            # Удаление устаревших таблиц
            tables_to_delete = SomeTables.objects.exclude(
                id__in=[table.id for table in processed_tables]
            )  # Исключаем обработанные
            deleted_count = tables_to_delete.count()
            tables_to_delete.delete()
            if deleted_count > 0:
                logger.info(f"Удалено {deleted_count} устаревших таблиц.")
        except Exception as e:
            logger.error(f"Ошибка при обработке раздела {e}")

        channel_layer = get_channel_layer()
        tables_id = SomeTables.objects.all()
        for num, table_zip in enumerate(zip(tables_id, tables)):
            progress: int = int((num / len(tables)) * 100)
            logger.info(f"Отправка прогресса: {progress}%")
            async_to_sync(channel_layer.group_send)(
                "progress_updates",
                {
                    "type": "send_progress",
                    "progress": progress,
                },
            )
            for row_num, row in enumerate(table_zip[1].rows[3:]):
                cells = [cell.text.strip().replace("\n", "<br>") for cell in row.cells]
                row_in_db = (
                    SomeDataFromSomeTables.objects.select_related("table_id")
                    .filter(table_id=table_zip[0].id, dock_num=row_num + 1)
                    .first()
                )

                value_corrector = {"+": True, "-": False}
                cls.model_inf: Dict[str, any] = {
                    "table_id": table_zip[0],
                    "dock_num": row_num + 1,
                    "location": cells[1],
                    "name_organ": cells[2],
                    "pseudonim": cells[3],
                    "letters": (
                        value_corrector.get(cells[4])
                        if value_corrector.get(cells[4])
                        else False
                    ),
                    "writing": (
                        value_corrector.get(cells[5])
                        if value_corrector.get(cells[5])
                        else False
                    ),
                    "ip_address": (
                        cells[6] if cells[6] not in ["+", "-"] else cells[7]
                    ),
                    "some_number": (
                        cells[7] if cells[6] not in ["+", "-"] else cells[8]
                    ),
                    "work_timme": (
                        cells[8] if cells[6] not in ["+", "-"] else cells[9]
                    ),
                }

                # cr_or_upd_row = SomeDataFromSomeTables(**cls.model_inf)
                # processed_cities.append(cr_or_upd_row)
                if row_in_db:
                    needs_update = False
                    for key, value in cls.model_inf.items():
                        if getattr(row_in_db, key) != value:
                            setattr(row_in_db, key, value)
                            needs_update = True

                    if needs_update:
                        cities_to_update.append(
                            row_in_db
                        )  # Добавляем в список для обновления только если были изменения
                    processed_cities.append(
                        row_in_db
                    )  # Добавляем в список обработанных в любом случае

                else:
                    cr_or_upd_row = SomeDataFromSomeTables(**cls.model_inf)
                    cities_to_add.append(cr_or_upd_row)
                    processed_cities.append(cr_or_upd_row)
            # if cr_or_upd_row[1]:
            #     logger.info(f"Добавил новую запись {str(cr_or_upd_row[0])}")

            # except Exception as e:
            #     logger.error(
            #         f"Ошибка при обработке записи Раздел {table_zip[0].table_name} -  Запись № {cells[0].text}"
            #     )
        try:
            if cities_to_add:
                res = SomeDataFromSomeTables.objects.bulk_create(cities_to_add)
                for row in res:
                    logger.info(f"Добавил новую запись id = '{row.id}' - {row.location}")
            if cities_to_update:
                SomeDataFromSomeTables.objects.bulk_update(
                    cities_to_update,
                    [
                        "location",
                        "name_organ",
                        "pseudonim",
                        "letters",
                        "writing",
                        "ip_address",
                        "some_number",
                        "work_timme",
                    ],
                )
                for el in cities_to_update:
                    logger.info(f"Обновил запись id = '{el.id}' - {el.location}")
            # Получите обновленный список городов
            delete_cities_ids = SomeDataFromSomeTables.objects.select_related(
                "table_id"
            ).exclude(id__in=[row.id for row in processed_cities])
            delete_cities_ids_count = delete_cities_ids.count()
            delete_cities_ids.delete()
            if delete_cities_ids_count > 0:
                logger.info(
                    f"Удалено {delete_cities_ids_count} устаревших записей городов."
                )
        except Exception as e:
            logger.error(f"Произошла ошибка при работе со строками таблицы {e}")

        all_rows = SomeDataFromSomeTables.objects.select_related("table_id").exclude(
            Q(location__isnull=True) | Q(location__exact=''))
        updated_cities = [
            el.to_dict() for el in all_rows
        ]  # Предполагается, что у вас есть метод to_dict()
        logger.info("Обработка завершена")

        # Отправьте финальное сообщение с прогрессом и обновленным списком городов
        async_to_sync(channel_layer.group_send)(
            "progress_updates",
            {
                "type": "send_progress",
                "progress": 100,
                "cities": updated_cities,  # Отправляем обновленный список городов
            },
        )

        #
        #

    @classmethod
    def style_cell_text(cls, cell, justifu: bool = False):
        """Стилизует текст в ячейке таблицы."""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justifu else WD_ALIGN_PARAGRAPH.CENTER  # Горизонтальное выравнивание

            # Установка стиля для каждого run в параграфе
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @classmethod
    def create_globus(cls):
        document = Document()

        # Получение секции документа
        section = document.sections[0]
        # Установка альбомной ориентации
        section.orientation = WD_ORIENT.LANDSCAPE
        # Явно задаем размеры страницы для альбомной ориентации
        section.page_width = Inches(11.69)  # Ширина для альбомной ориентации (A4)
        section.page_height = Inches(8.27)  # Высота для альбомной ориентации (A4)
        top_margin_cm = 2
        bottom_margin_cm = 2
        left_margin_cm = 3
        right_margin_cm = 1.5
        # Преобразование сантиметров в дюймы и установка полей
        section.top_margin = Inches(top_margin_cm / 2.54)
        section.bottom_margin = Inches(bottom_margin_cm / 2.54)
        section.left_margin = Inches(left_margin_cm / 2.54)
        section.right_margin = Inches(right_margin_cm / 2.54)
        # Стиль для заголовка
        style1 = document.styles.add_style('CustomHeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
        style1.base_style = document.styles['Heading 1']  # Наследуем от Heading 1
        font1 = style1.font
        font1.name = 'Times New Roman'
        font1._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')  # Для правильной установки шрифта
        font1.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        font1.size = Pt(16)

        paragraph_format = style1.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Одинарный интервал
        paragraph_format.space_before = Pt(0)  # Убираем отступ перед абзацем
        paragraph_format.space_after = Pt(0)  # Убираем отступ после абзаца
        paragraph_format.first_line_indent = Inches(0)  # Убираем отступ первой строки

        tables_name = SomeTables.objects.values_list('id', 'table_name')
        table_headers = ['№ №', 'Место положения', 'Наименование органа', 'Название пункта\n(псевдоним)',
                         'Вид передаваемой информации', 'Какой-то номер', 'Время работы, телефон для связи']
        table_headers2 = ['Письма', 'Написание', 'Сетевой адрес IP']
        for num, name in enumerate(tables_name):
            paragraph = document.add_paragraph(name[1], style='CustomHeadingStyle')
            document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Выравнивание по центру

            table_data = SomeDataFromSomeTables.objects.filter(table_id=name[0]).values_list(
                'table_id', 'dock_num', 'location', 'name_organ', 'pseudonim', 'letters', 'writing', 'ip_address',
                'some_number', 'work_timme')
            # Добавление таблицы
            table = document.add_table(rows=3 + len(table_data), cols=9)
            table.cell(0, 0).merge(table.cell(2, 0))  # Объединяем ячейки
            table.cell(0, 1).merge(table.cell(2, 1))  # Объединяем ячейки
            table.cell(0, 2).merge(table.cell(2, 2))  # Объединяем ячейки
            table.cell(0, 3).merge(table.cell(2, 3))  # Объединяем ячейки
            #
            table.cell(0, 4).merge(table.cell(0, 6))  # Объединяем ячейки
            table.cell(1, 4).merge(table.cell(2, 4))  # Объединяем ячейки
            table.cell(1, 5).merge(table.cell(2, 5))  # Объединяем ячейки
            table.cell(1, 6).merge(table.cell(2, 6))  # Объединяем ячейки

            table.cell(0, 7).merge(table.cell(2, 7))  # Объединяем ячейки
            table.cell(0, 8).merge(table.cell(2, 8))  # Объединяем ячейки

            table.style = 'Table Grid'  # Установка стиля таблицы с границами

            # Добавление заголовков таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = table_headers[0]
            hdr_cells[1].text = table_headers[1]
            hdr_cells[2].text = table_headers[2]
            hdr_cells[3].text = table_headers[3]
            hdr_cells[4].text = table_headers[4]
            hdr_cells[7].text = table_headers[5]
            hdr_cells[8].text = table_headers[6]
            hdr_cells1 = table.rows[1].cells
            for i, header in enumerate(table_headers2):
                hdr_cells1[i + 4].text = header
            # Установка стиля для всех ячеек
            for row_num, row in enumerate(table.rows):
                for cell_num, cell in enumerate(row.cells):
                    cls.style_cell_text(cell)
                    if row_num >= 3:
                        # Получаем данные для вставки
                        data_insert = str(table_data[row_num - 3][cell_num + 1]).replace('<br>', '\n')
                        if data_insert == 'True':
                            data_insert = '+'
                        elif data_insert == 'False':
                            data_insert = '-'
                        cell.text = data_insert
                        if cell_num in [0, 1, 3, 4, 5, 7]:
                            cls.style_cell_text(cell)
                        else:
                            cls.style_cell_text(cell, justifu=True)
            document.add_page_break()
            document.save(os.path.join(ProjectSettings.tlg_dir, 'globus_new.docx'))


if __name__ == "__main__":
    # GlobusParser.process_file("globus.docx")
    GlobusParser.create_globus()
