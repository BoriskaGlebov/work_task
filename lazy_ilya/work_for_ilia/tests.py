import json
import os
from unittest.mock import patch, MagicMock

from django.contrib.auth.models import User
from django.test import TestCase, Client
from django.urls import reverse
from docx import Document

from work_for_ilia.models import Counter
from .utils.my_settings.disrs_for_app import ProjectSettings, logger


class GreaterViewTests(TestCase):
    """
    Тесты для представления Greater.
    """

    @classmethod
    def setUpClass(cls):
        cls.test_file_path = 'test_file.doc'
        doc = Document()
        cls.test_str = 'Test content for .doc file.'
        doc.add_paragraph(cls.test_str)
        doc.save(cls.test_file_path)

    def setUp(self) -> None:
        """
        Настройка тестового клиента и URL-адресов для тестов.
        """
        self.client: Client = Client()
        self.upload_url: str = reverse('work_for_ilia:index')
        self.update_url: str = reverse('work_for_ilia:update')

    def test_get_index_view(self) -> None:
        """
        Тестирует успешное получение страницы загрузки.
        Из-за моков тест падает, но если отдельно запускать то нормально работает.
        НАдо так проверить, что файл реально верно сохраняется.
        """
        response = self.client.get(self.upload_url)
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'work_for_ilia/index.html')

    @patch('work_for_ilia.views.Converter')
    @patch('work_for_ilia.views.Parser')  # Мокируем Parser, если он тоже используется
    def test_post_upload_files_success(self, mock_parser, mock_converter) -> None:
        """
        Тестирует успешную загрузку файлов .doc.
        """
        # Настройка мока для Converter
        mock_converter.return_value.convert_files.return_value = None  # Имитация успешной конвертации

        # Настройка мока для Parser
        mock_parser.return_value.create_file_parsed.return_value = self.test_str  # Имитация возвращаемого содержимого

        with open(self.test_file_path, 'rb') as f:
            response = self.client.post(self.upload_url, {
                'file': [f],
                'document_number': 1,
            })

        self.assertEqual(response.status_code, 200)
        json_response = json.loads(response.content)

        # Проверяем наличие ключа 'content' и его значение
        self.assertIn('content', json_response)
        self.assertEqual(json_response['content'],
                         self.test_str)  # Проверяем, что содержимое соответствует ожиданиям

    def test_post_upload_files_success2(self) -> None:
        """
        Тестирует успешную загрузку файлов .doc.
        """

        # Открываем файл для загрузки
        with open(self.test_file_path, 'rb') as f:
            response = self.client.post(self.upload_url, {
                'file': [f, ],
                'document_number': 1,
            })

        # Проверяем статус ответа
        self.assertEqual(response.status_code, 200)

        # Декодируем JSON-ответ
        json_response = json.loads(response.content)
        logger.info(f'Содержание json ответа от сервера\n{json_response}')
        # Проверяем наличие ожидаемых ключей в ответе
        self.assertIn('content', json_response)
        self.assertIn('new_files', json_response)

        # Проверяем, что новое имя файла соответствует ожиданиям
        expected_new_file_name = '1_test_file.txt'  # Замените на ожидаемое имя файла после обработки
        self.assertIn(expected_new_file_name, json_response['new_files'])
        self.assertIn(self.test_str.upper(), json_response['content'][0])

    #
    def test_post_upload_files_no_files(self) -> None:
        """
        Тестирует ситуацию, когда файлы не загружены.
        """
        response = self.client.post(self.upload_url, {
            'document_number': 2,
        })

        self.assertEqual(response.status_code, 400)
        json_response = json.loads(response.content)
        self.assertEqual(json_response['error'], 'Нет загруженных файлов')

    def test_post_upload_files_invalid_document_number(self) -> None:
        """
        Тестирует ситуацию с неверным номером документа.
        """

        with open(self.test_file_path, 'rb') as f:
            response = self.client.post(self.upload_url, {
                'file': [f, ],  # Пустой список файлов
                'document_number': 0,  # Неверный номер документа
            })

        self.assertEqual(response.status_code, 400)
        json_response = json.loads(response.content)
        self.assertEqual(json_response['error'], 'Номер документа должен быть больше нуля')

    def test_post_upload_files_file_save_error(self) -> None:
        """
        Тестирует ситуацию с ошибкой при сохранении файла.
        """
        # Здесь мы можем использовать мок для имитации ошибки сохранения файла.
        # Например, можно использовать unittest.mock для замены метода save.

        from unittest.mock import patch

        with patch('work_for_ilia.views.OverwritingFileSystemStorage.save',
                   side_effect=Exception("Ошибка сохранения файла")):
            with open(self.test_file_path, 'rb') as f:
                response = self.client.post(self.upload_url, {
                    'file': [f, ],  # Имитация загрузки файла
                    'document_number': 3,
                })

            self.assertEqual(response.status_code, 500)
            json_response = json.loads(response.content)
            self.assertIn('error', json_response)

    @patch("work_for_ilia.views.Converter")
    def test_post_upload_files_processing_error(self, mock_converter) -> None:
        """
        Тестирует ситуацию с ошибкой при обработке файла.
        """
        mock_converter.side_effect = Exception("Ошибка обработки файла")
        with open(self.test_file_path, 'rb') as f:
            response = self.client.post(self.upload_url, {
                'file': [f, ],  # Имитация загрузки файла
                'document_number': 4,
            })

        self.assertEqual(response.status_code, 500)
        json_response = json.loads(response.content)
        self.assertIn('error', json_response)

    def test_put_update_files_success(self) -> None:
        """
        Тестирует успешное обновление файлов.
        """
        # Предположим, что у нас есть уже загруженные файлы
        Counter.objects.create(num_files=5)
        # Открываем файл для загрузки
        with open(self.test_file_path, 'rb') as f:
            response = self.client.post(self.upload_url, {
                'file': [f, ],
                'document_number': 5,
            })

        response = self.client.put(self.update_url, json.dumps([
            {
                'document_number': 5,
                'content': 'Updated content',
                'new_file_name': '5_test_file.txt'
            }
        ]), content_type='application/json')

        self.assertEqual(response.status_code, 200)
        json_response = json.loads(response.content)
        self.assertEqual(json_response['status'], 'success')
        os.remove(os.path.join(ProjectSettings.tlg_dir, '5_test_file.txt'))

    @classmethod
    def tearDownClass(cls):
        os.remove(cls.test_file_path)
        # pass


class CitiesTests(TestCase):
    def setUp(self):
        self.client = Client()
        self.url = reverse('work_for_ilia:cities')  # Замените 'cities' на имя вашего URL
        self.superuser = self.create_superuser()  # Создаем суперпользователя для тестирования

    def create_superuser(self):
        return User.objects.create_superuser(
            username='admin',
            email='admin@example.com',
            password='adminpass'
        )

    def test_post_success(self):
        self.client.force_login(self.superuser)  # Логинимся как суперпользователь

        with patch('work_for_ilia.views.OverwritingFileSystemStorage') as mock_fs:
            mock_file = MagicMock(name='file', spec='File')
            mock_file.name = 'globus.docx'  # Убедитесь, что имя файла соответствует ожидаемому
            mock_fs.return_value.save.return_value = 'mock_file_path'

            response = self.client.post(self.url, {
                'cityFile': mock_file,
            })

            self.assertEqual(response.status_code, 200)
            self.assertJSONEqual(response.content, {'message': 'File uploaded successfully'})

    def test_post_no_file(self):
        self.client.login(username='admin', password='adminpass')  # Логинимся как суперпользователь

        response = self.client.post(self.url, {})

        self.assertEqual(response.status_code, 400)
        self.assertJSONEqual(response.content, {'error': 'No file uploaded'})

    @patch('work_for_ilia.views.Document')  # Мокируем Document для обработки файла
    @patch('work_for_ilia.views.SomeTables.objects.bulk_create')  # Мокируем bulk_create
    def test_process_file_success(self, mock_bulk_create, mock_document):
        self.client.login(username='admin', password='adminpass')  # Логинимся как суперпользователь

        mock_document.return_value.tables = []  # Мокируем таблицы документа
        mock_document.return_value.paragraphs = [MagicMock(text='City1'), MagicMock(text='City2')]

        # Запускаем обработку файла
        Cities.process_file('mock_file_path')

        # Проверяем, что новые таблицы были созданы правильно
        self.assertEqual(mock_bulk_create.call_count, 1)

    def test_get(self):
        self.client.login(username='admin', password='adminpass')  # Логинимся как суперпользователь

        response = self.client.get(self.url)

        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'work_for_ilia/cities.html')

    def test_get_cities_json(self):
        self.client.login(username='admin', password='adminpass')  # Логинимся как суперпользователь

        SomeTables.objects.create(table_name='Moscow')  # Создаем тестовые данные

        response = self.client.get(self.url)

        self.assertEqual(response.status_code, 200)

        cities_json = response.context['cities_json']
        expected_data = [{'table_name': 'Moscow'}]  # Предполагается, что у вас есть метод to_dict() в модели

        self.assertJSONEqual(cities_json, expected_data)


class StatisticViewTests(TestCase):
    """
    Тесты для представления Statistic.
    """

    def setUp(self) -> None:
        """
        Настройка тестового клиента и URL-адреса для тестов.
        """
        self.client: Client = Client()
        self.statistics_url: str = reverse('work_for_ilia:statistics')

    def test_get_statistics_view(self) -> None:
        """
        Тестирует успешное получение страницы статистики.
        """
        # Создаем несколько записей в Counter для тестирования
        Counter.objects.create(num_files=10)

        response = self.client.get(self.statistics_url)

        # Проверяем статус ответа и используемую шаблон
        self.assertEqual(response.status_code, 200)
        self.assertTemplateUsed(response, 'work_for_ilia/statistics.html')

        # Проверяем контекст
        context = response.context
        self.assertIn('converted_files', context)
        self.assertEqual('10', context['converted_files'])
