import os
from pathlib import Path

from docx.table import _Cell, Table

# Укажите путь к настройкам вашего проекта
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "lazy_ilya.settings")
import django

# Настройка Django
django.setup()
# import asyncio

import re
from pprint import pprint
from queue import Queue

from django.db.models import Q, QuerySet
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.font import Font

from cities.models import TableNames, CityData

from typing import Any, Callable, Dict, List, Optional, Tuple

from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer
from docx import Document

from lazy_ilya.utils.settings_for_app import (ProjectSettings, logger)


class GlobusParser:
    """
    Класс обрабатывает содержимое файла с Городами.
    """

    def __init__(self) -> None:
        """
        Инициализирует экземпляр класса GlobusParser.

        Создает структуру для хранения информации о городах.
        """
        self.model_inf: Dict[str, Any] = {
            "table_id": "",
            "dock_num": "",
            "location": "",
            "name_organ": "",
            "pseudonim": "",
            "letters": "",
            "writing": "",
            "ip_address": "",
            "some_number": "",
            "work_time": "",
        }

    # @classmethod
    # @logger.catch(message="Непредвиденное исключение")
    # def process_file(cls, file_path: str) -> None:
    #     """
    #     Обрабатывает загруженный файл с данными о городах.
    #     Обновляет названия таблиц на основе номера раздела.
    #
    #     Args:
    #         file_path (str): Путь к загруженному файлу.
    #     """
    #     doc: Document = Document(ProjectSettings.tlg_dir / file_path)
    #     tables: List[Any] = doc.tables
    #     paragraphs: List[Any] = doc.paragraphs
    #     processed_tables: List[TableNames] = []  # Список обработанных разделов
    #     tables_to_add: List[TableNames] = []
    #     tables_to_update: List[TableNames] = []
    #
    #     processed_cities: List[TableNames] = []
    #     cities_to_add: List[TableNames] = []
    #     cities_to_update: List[TableNames] = []
    #
    #     for paragraph in paragraphs[1:]:
    #         pattern: str = r"^Раздел (\d+)\s*(.*)"  # Группировка номера и названия
    #         text: str = paragraph.text.strip()
    #
    #         match: re.Match = re.match(pattern, text)
    #
    #         if match:
    #             section_number: str = match.group(1)  # Номер раздела
    #             section_name: str = text  # Полное название раздела
    #             # Поиск таблицы по номеру раздела (в table_name)
    #             table: Optional[TableNames] = TableNames.objects.filter(
    #                 Q(table_name__startswith=f"Раздел {section_number}")
    #                 & Q(table_name__regex=rf"^Раздел {section_number}(?!\d)")
    #             ).first()  # Ищем по началу строки
    #             if table:
    #                 # Обновление существующей таблицы
    #                 if table.table_name != section_name:
    #                     table.table_name = section_name
    #                     tables_to_update.append(table)
    #                     processed_tables.append(table)  # Запоминаем раздел
    #                 elif table and table.table_name == section_name:
    #                     processed_tables.append(table)  # Запоминаем раздел
    #             else:
    #                 # Создание новой таблицы
    #                 table = TableNames(table_name=section_name)
    #                 tables_to_add.append(table)
    #                 processed_tables.append(table)  # Запоминаем ID
    #     try:
    #         if tables_to_add:
    #             res: List[TableNames] = TableNames.objects.bulk_create(tables_to_add)
    #             for el in res:
    #                 logger.info(
    #                     f"Создана новая таблица '{el.table_name}' (ID: {el.id})"
    #                 )
    #         if tables_to_update:
    #             res2: int = TableNames.objects.bulk_update(
    #                 tables_to_update,
    #                 [
    #                     "table_name",
    #                 ],
    #             )
    #             for el in tables_to_update:
    #                 logger.info(
    #                     f"Обновлено название таблицы '{el.table_name}' (ID: {el.id})"
    #                 )
    #         # Удаление устаревших таблиц
    #         tables_to_delete = TableNames.objects.exclude(
    #             id__in=[table.id for table in processed_tables]
    #         )  # Исключаем обработанные
    #         deleted_count: int = tables_to_delete.count()
    #         tables_to_delete.delete()
    #         if deleted_count > 0:
    #             logger.info(f"Удалено {deleted_count} устаревших таблиц.")
    #     except Exception as e:
    #         logger.error(f"Ошибка при обработке раздела {e}")
    #
    #     channel_layer = get_channel_layer()
    #     tables_id = TableNames.objects.all()
    #     for num, table_zip in enumerate(zip(tables_id, tables)):
    #         progress: int = int((num / len(tables)) * 100)
    #         logger.info(f"Отправка прогресса: {progress}%")
    #         async_to_sync(channel_layer.group_send)(
    #             "progress_updates",
    #             {
    #                 "type": "send_progress",
    #                 "progress": progress,
    #             },
    #         )
    #         for row_num, row in enumerate(table_zip[1].rows[3:]):
    #             cells: List[str] = [
    #                 cell.text.strip().replace("\n", "<br>") for cell in row.cells
    #             ]
    #             row_in_db: Optional[CityData] = (
    #                 CityData.objects.select_related("table_id")
    #                 .filter(table_id=table_zip[0].id, dock_num=row_num + 1)
    #                 .first()
    #             )
    #
    #             value_corrector: Dict[str, bool] = {"+": True, "-": False}
    #             cls.model_inf: Dict[str, Any] = {
    #                 "table_id": table_zip[0],
    #                 "dock_num": row_num + 1,
    #                 "location": cells[1],
    #                 "name_organ": cells[2],
    #                 "pseudonim": cells[3],
    #                 "letters": (
    #                     value_corrector.get(cells[4])
    #                     if value_corrector.get(cells[4])
    #                     else False
    #                 ),
    #                 "writing": (
    #                     value_corrector.get(cells[5])
    #                     if value_corrector.get(cells[5])
    #                     else False
    #                 ),
    #                 "ip_address": (
    #                     cells[6] if cells[6] not in ["+", "-"] else cells[7]
    #                 ),
    #                 "some_number": (
    #                     cells[7] if cells[6] not in ["+", "-"] else cells[8]
    #                 ),
    #                 "work_timme": (
    #                     cells[8] if cells[6] not in ["+", "-"] else cells[9]
    #                 ),
    #             }
    #             if row_in_db:
    #                 needs_update: bool = False
    #                 for key, value in cls.model_inf.items():
    #                     if getattr(row_in_db, key) != value:
    #                         setattr(row_in_db, key, value)
    #                         needs_update = True
    #
    #                 if needs_update:
    #                     cities_to_update.append(
    #                         row_in_db
    #                     )  # Добавляем в список для обновления только если были изменения
    #                 processed_cities.append(
    #                     row_in_db
    #                 )  # Добавляем в список обработанных в любом случае
    #
    #             else:
    #                 cr_or_upd_row: CityData = CityData(
    #                     **cls.model_inf
    #                 )
    #                 cities_to_add.append(cr_or_upd_row)
    #                 processed_cities.append(cr_or_upd_row)
    #     try:
    #         if cities_to_add:
    #             res: List[CityData] = (
    #                 CityData.objects.bulk_create(cities_to_add)
    #             )
    #             for row in res:
    #                 logger.info(
    #                     f"Добавил новую запись id = '{row.id}' - {row.location}"
    #                 )
    #         if cities_to_update:
    #             CityData.objects.bulk_update(
    #                 cities_to_update,
    #                 [
    #                     "location",
    #                     "name_organ",
    #                     "pseudonim",
    #                     "letters",
    #                     "writing",
    #                     "ip_address",
    #                     "some_number",
    #                     "work_timme",
    #                 ],
    #             )
    #             for el in cities_to_update:
    #                 logger.info(f"Обновил запись id = '{el.id}' - {el.location}")
    #         # Получите обновленный список городов
    #         delete_cities_ids = CityData.objects.select_related(
    #             "table_id"
    #         ).exclude(id__in=[row.id for row in processed_cities])
    #         delete_cities_ids_count: int = delete_cities_ids.count()
    #         delete_cities_ids.delete()
    #         if delete_cities_ids_count > 0:
    #             logger.info(
    #                 f"Удалено {delete_cities_ids_count} устаревших записей городов."
    #             )
    #     except Exception as e:
    #         logger.error(f"Произошла ошибка при работе со строками таблицы {e}")
    #
    #     all_rows = CityData.objects.select_related("table_id").exclude(
    #         Q(location__isnull=True) | Q(location__exact="")
    #     )
    #     updated_cities: List[Dict[str, Any]] = [
    #         el.to_dict() for el in all_rows
    #     ]  # Предполагается, что у вас есть метод to_dict()
    #     logger.info("Обработка завершена")
    #
    #     # Отправьте финальное сообщение с прогрессом и обновленным списком городов
    #     async_to_sync(channel_layer.group_send)(
    #         "progress_updates",
    #         {
    #             "type": "send_progress",
    #             "progress": 100,
    #             "cities": updated_cities,  # Отправляем обновленный список городов
    #         },
    #     )
    @classmethod
    def process_file(cls, file_path: str) -> None:
        """
        Обрабатывает документ по указанному пути.

        Параметры:
        -----------
        file_path : str
            Путь к файлу документа для обработки.

        Логика работы:
        ---------------
        1. Загружает документ с помощью метода _load_doc.
        2. Извлекает таблицы и параграфы из документа.
        3. Обрабатывает параграфы методом _process_paragraphs, который возвращает:
            - processed_tables: список обработанных таблиц (структура не уточнена, возможно List[Any])
            - tables_to_add: таблицы, которые необходимо добавить
            - tables_to_update: таблицы, которые необходимо обновить
        4. Синхронизирует таблицы через _sync_tables.
        5. Обрабатывает таблицы с учётом строк через _process_tables_with_rows.
        6. Финализирует прогресс обработки через _finalize_progress.

        Исключения:
        -----------
        Любые исключения логируются с уровнем ошибки.

        Возвращаемое значение:
        ----------------------
        None
        """
        try:
            doc = cls._load_doc(
                file_path)  # тип doc зависит от реализации _load_doc (вероятно, объект python-docx.Document)
            tables = doc.tables  # Обычно список объектов таблиц документа
            paragraphs = doc.paragraphs  # Список параграфов документа

            # Предполагаемые типы возвращаемых значений:
            # processed_tables: List[Any]
            # tables_to_add: List[Any]
            # tables_to_update: List[Any]
            processed_tables, tables_to_add, tables_to_update = cls._process_paragraphs(paragraphs)

            # Синхронизация таблиц: метод без возвращаемого значения
            cls._sync_tables(processed_tables, tables_to_add, tables_to_update)

            # Обработка таблиц с учетом строк: метод без возвращаемого значения
            cls._process_tables_with_rows(tables, processed_tables)

            # Завершение обработки
            cls._finalize_progress()

        except Exception as e:
            logger.exception(f"Непредвиденное исключение: {e}")

    @staticmethod
    def _load_doc(file_path: str) -> Document:
        """
        Загружает документ Word из указанного пути.

        Параметры:
        -----------
        file_path : str
            Путь к файлу документа.

        Возвращаемое значение:
        ----------------------
        Document
            Объект документа Word (из python-docx).
        """
        return Document(ProjectSettings.tlg_dir / file_path)

    @classmethod
    def _process_paragraphs(cls, paragraphs: List[Any]) -> Tuple[
        List["TableNames"], List["TableNames"], List["TableNames"]]:
        """
        Обрабатывает список параграфов, выделяя из них разделы и соответствующие таблицы.

        Параметры:
        -----------
        paragraphs : List[Any]
            Список параграфов документа (объекты из python-docx или подобные).

        Логика:
        --------
        - Для каждого параграфа (кроме первого) проверяется, соответствует ли он шаблону "Раздел N ...".
        - Если таблица с таким именем существует в базе, проверяет совпадение имени:
          - Если имя изменилось — добавляет таблицу в список для обновления.
          - Иначе — просто добавляет в обработанные.
        - Если таблицы нет — создает новый объект для добавления.

        Возвращаемое значение:
        ----------------------
        Tuple[
            List[TableNames],  # processed_tables — таблицы, учтённые для дальнейшей обработки
            List[TableNames],  # tables_to_add — новые таблицы для добавления в БД
            List[TableNames]   # tables_to_update — существующие таблицы с изменёнными именами
        ]
        """
        processed_tables, tables_to_add, tables_to_update = [], [], []

        for paragraph in paragraphs[1:]:
            match = re.match(r"^Раздел (\d+)\s*(.*)", paragraph.text.strip())
            if match:
                section_number = match.group(1)
                section_name = paragraph.text.strip()
                table = TableNames.objects.filter(
                    Q(table_name__startswith=f"Раздел {section_number}") &
                    Q(table_name__regex=rf"^Раздел {section_number}(?!\d)")
                ).first()

                if table:
                    if table.table_name != section_name:
                        table.table_name = section_name
                        tables_to_update.append(table)
                    processed_tables.append(table)
                else:
                    new_table = TableNames(table_name=section_name)
                    tables_to_add.append(new_table)
                    processed_tables.append(new_table)

        return processed_tables, tables_to_add, tables_to_update

    @classmethod
    def _sync_tables(cls,
                     processed_tables: List["TableNames"],
                     tables_to_add: List["TableNames"],
                     tables_to_update: List["TableNames"]) -> None:
        """
        Синхронизирует таблицы с базой данных: добавляет новые, обновляет изменённые и удаляет устаревшие.

        Параметры:
        -----------
        processed_tables : List[TableNames]
            Список таблиц, которые должны остаться в базе (актуальные).

        tables_to_add : List[TableNames]
            Список новых таблиц для добавления в базу.

        tables_to_update : List[TableNames]
            Список таблиц, требующих обновления (изменение названия).

        Возвращаемое значение:
        ----------------------
        None
        """
        if tables_to_add:
            created = TableNames.objects.bulk_create(tables_to_add)
            for t in created:
                logger.info(f"Создана новая таблица '{t.table_name}' (ID: {t.id})")

        if tables_to_update:
            TableNames.objects.bulk_update(tables_to_update, ["table_name"])
            for t in tables_to_update:
                logger.info(f"Обновлено название таблицы '{t.table_name}' (ID: {t.id})")

        deleted = TableNames.objects.exclude(id__in=[t.id for t in processed_tables])
        deleted_count = deleted.count()
        deleted.delete()
        if deleted_count:
            logger.info(f"Удалено {deleted_count} устаревших таблиц.")

    @classmethod
    def _process_tables_with_rows(cls,
                                  tables: List[Any],
                                  tables_id: List["TableNames"]) -> None:
        """
        Обрабатывает строки таблиц из документа, синхронизирует данные с базой.

        Параметры:
        -----------
        tables : List[Any]
            Список таблиц из документа (например, объекты таблиц python-docx).

        tables_id : List[TableNames]
            Список моделей таблиц из базы данных, соответствующих таблицам документа.

        Логика:
        --------
        - Для каждой таблицы (параллельно с моделью таблицы) проходит по строкам,
          начиная с четвертой (index 3).
        - Извлекает и корректирует данные из ячеек.
        - Ищет запись в базе по table_id и dock_num (номер строки).
        - Если запись существует, обновляет при необходимости.
        - Если нет — создает новую.
        - Отправляет прогресс обработки через Channels.
        - В конце вызывает синхронизацию данных с БД (_sync_city_data).

        Возвращаемое значение:
        ----------------------
        None
        """
        channel_layer = get_channel_layer()
        processed_cities, cities_to_add, cities_to_update = [], [], []

        for num, (table_model, doc_table) in enumerate(zip(tables_id, tables)):
            progress = int((num / len(tables)) * 100)
            logger.info(f"Отправка прогресса: {progress}%")
            async_to_sync(channel_layer.group_send)(
                "progress_updates", {"type": "send_progress", "progress": progress}
            )

            for row_num, row in enumerate(doc_table.rows[3:]):
                # cells = [cell.text.strip().replace("\n", "<br>") for cell in row.cells]
                cells = [cell.text.strip() for cell in row.cells]
                row_in_db = CityData.objects.select_related("table_id").filter(
                    table_id=table_model.id,
                    dock_num=row_num + 1
                ).first()

                value_corrector = {"+": True, "-": False}
                cls.model_inf = {
                    "table_id": table_model,
                    "dock_num": row_num + 1,
                    "location": cells[1],
                    "name_organ": cells[2],
                    "pseudonim": cells[3],
                    "letters": value_corrector.get(cells[4], False),
                    "writing": value_corrector.get(cells[5], False),
                    "ip_address": cells[6] if cells[6] not in ["+", "-"] else cells[7],
                    "some_number": cells[7] if cells[6] not in ["+", "-"] else cells[8],
                    "work_time": cells[8] if cells[6] not in ["+", "-"] else cells[9],
                }

                if row_in_db:
                    needs_update = False
                    for key, value in cls.model_inf.items():
                        if getattr(row_in_db, key) != value:
                            setattr(row_in_db, key, value)
                            needs_update = True
                    if needs_update:
                        cities_to_update.append(row_in_db)
                    processed_cities.append(row_in_db)
                else:
                    new_row = CityData(**cls.model_inf)
                    cities_to_add.append(new_row)
                    processed_cities.append(new_row)

        cls._sync_city_data(cities_to_add, cities_to_update, processed_cities)

    @classmethod
    def _sync_city_data(cls,
                        cities_to_add: List["CityData"],
                        cities_to_update: List["CityData"],
                        processed_cities: List["CityData"]) -> None:
        """
        Синхронизирует записи городов с базой: добавляет новые, обновляет существующие и удаляет устаревшие.

        Параметры:
        -----------
        cities_to_add : List[CityData]
            Список новых записей для добавления.

        cities_to_update : List[CityData]
            Список записей для обновления.

        processed_cities : List[CityData]
            Все записи, которые должны остаться в базе.

        Возвращаемое значение:
        ----------------------
        None
        """
        if cities_to_add:
            created = CityData.objects.bulk_create(cities_to_add)
            for row in created:
                logger.info(f"Добавил новую запись id = '{row.id}' - {row.location}")

        if cities_to_update:
            CityData.objects.bulk_update(cities_to_update, [
                "location", "name_organ", "pseudonim", "letters",
                "writing", "ip_address", "some_number", "work_time"
            ])
            for row in cities_to_update:
                logger.info(f"Обновил запись id = '{row.id}' - {row.location}")

        to_delete = CityData.objects.select_related("table_id").exclude(
            id__in=[r.id for r in processed_cities]
        )
        deleted_count = to_delete.count()
        to_delete.delete()
        if deleted_count:
            logger.info(f"Удалено {deleted_count} устаревших записей городов.")

    @classmethod
    def _finalize_progress(cls) -> None:
        """
        Завершает процесс обновления, отправляя в канал итоговые данные.

        Логика:
        --------
        - Получает все актуальные записи городов с непустым location.
        - Преобразует их в словари (через метод to_dict модели).
        - Отправляет сообщение с прогрессом 100% и списком городов через Channels.

        Возвращаемое значение:
        ----------------------
        None
        """
        all_rows = CityData.objects.select_related("table_id").exclude(
            Q(location__isnull=True) | Q(location__exact="")
        )
        updated_cities = [el.to_dict() for el in all_rows]
        channel_layer = get_channel_layer()
        async_to_sync(channel_layer.group_send)(
            "progress_updates",
            {
                "type": "send_progress",
                "progress": 100,
                "cities": updated_cities
            }
        )
        logger.info(f"Отправка прогресса: 100%")

    @classmethod
    def _style_cell_text(cls, cell: _Cell, justify: bool = False) -> None:
        """
        Применяет стиль к тексту внутри ячейки таблицы Word.

        Устанавливает горизонтальное выравнивание абзацев, вертикальное выравнивание ячейки,
        шрифт Times New Roman и размер шрифта 12 пунктов для всех текстовых фрагментов (runs).

        Args:
            cell (_Cell): Объект ячейки таблицы из библиотеки python-docx.
            justify (bool, optional): Если True, выравнивание текста по ширине (Justify),
                                      иначе по центру. По умолчанию False.

        Returns:
            None
        """
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.CENTER
            )

            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    @classmethod
    def _init_document(cls) -> Document:
        """
        Создает и настраивает новый документ Word с альбомной ориентацией и заданными полями.

        Returns:
            Document: Объект документа с установленной альбомной ориентацией и полями.
        """
        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)  # Ширина A4 в альбомной ориентации
        section.page_height = Inches(8.27)  # Высота A4 в альбомной ориентации
        section.top_margin = Inches(2 / 2.54)  # 2 см
        section.bottom_margin = Inches(2 / 2.54)  # 2 см
        section.left_margin = Inches(3 / 2.54)  # 3 см
        section.right_margin = Inches(1.5 / 2.54)  # 1.5 см
        return document

    @classmethod
    def _create_custom_heading_style(cls, document: Document) -> Any:
        """
        Создает и возвращает пользовательский стиль заголовка для документа.

        Args:
            document (Document): Объект документа, в котором создается стиль.

        Returns:
            _ParagraphStyle: Объект созданного стиля абзаца с заданными параметрами.
        """
        style = document.styles.add_style("CustomHeadingStyle", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Heading 1"]  # Наследование от Heading 1
        font = style.font
        font.name = "Times New Roman"
        font._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")  # Корректная установка шрифта для ASCII
        font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
        font.size = Pt(16)  # Размер шрифта 16 пунктов

        para_fmt = style.paragraph_format
        para_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
        para_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE  # Одинарный межстрочный интервал
        para_fmt.space_before = Pt(0)  # Отступ перед абзацем
        para_fmt.space_after = Pt(0)  # Отступ после абзаца
        para_fmt.first_line_indent = Inches(0)  # Отступ первой строки

        return style

    @classmethod
    def _create_and_format_table(cls, document: Document, num_rows: int) -> Table:
        """
        Создает таблицу с фиксированным числом столбцов (9) и заданным количеством строк,
        а также выполняет объединение ячеек по заданной схеме.

        Args:
            document (Document): Объект документа, куда будет добавлена таблица.
            num_rows (int): Количество строк в таблице.

        Returns:
            Table: Объект таблицы с примененными объединениями и стилем.
        """
        table = document.add_table(rows=num_rows, cols=9)

        # Объединения ячеек
        table.cell(0, 0).merge(table.cell(2, 0))
        table.cell(0, 1).merge(table.cell(2, 1))
        table.cell(0, 2).merge(table.cell(2, 2))
        table.cell(0, 3).merge(table.cell(2, 3))
        table.cell(0, 4).merge(table.cell(0, 6))
        table.cell(1, 4).merge(table.cell(2, 4))
        table.cell(1, 5).merge(table.cell(2, 5))
        table.cell(1, 6).merge(table.cell(2, 6))
        table.cell(0, 7).merge(table.cell(2, 7))
        table.cell(0, 8).merge(table.cell(2, 8))

        table.style = "Table Grid"  # Стиль с границами

        return table

    @classmethod
    def _fill_table_headers(cls, table: Table) -> None:
        """
        Заполняет заголовки таблицы в первых двух строках.

        Args:
            table (Table): Таблица Word, в которую нужно вставить заголовки.
        """
        headers1 = [
            "№ №",
            "Место положения",
            "Наименование органа",
            "Название пункта\n(псевдоним)",
            "Вид передаваемой информации",
            "", "",  # колонки 5 и 6 объединены с 4
            "Какой-то номер",
            "Время работы, телефон для связи",
        ]
        headers2 = ["Письма", "Написание", "Сетевой адрес IP"]

        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers1):
            hdr_cells[i].text = text

        hdr_cells2 = table.rows[1].cells
        for i, text in enumerate(headers2):
            hdr_cells2[i + 4].text = text

    @classmethod
    def _fill_table_data(cls, table: Table, table_data: List[Tuple]) -> None:
        """
        Заполняет таблицу данными из списка кортежей.
        Применяет стили к ячейкам.

        Args:
            table (Table): Таблица для заполнения.
            table_data (List[Tuple]): Данные для заполнения таблицы. Каждый кортеж — строка данных.
        """
        for row_num, row in enumerate(table.rows):
            for cell_num, cell in enumerate(row.cells):
                cls._style_cell_text(cell)
                if row_num >= 3:
                    # Извлекаем данные, сдвигаем индексы под таблицу (пропускаем 3 заголовка)
                    data_insert = str(table_data[row_num - 3][cell_num + 1]).replace("<br>", "\n")
                    if data_insert == "True":
                        data_insert = "+"
                    elif data_insert == "False":
                        data_insert = "-"
                    cell.text = data_insert

                    # Применяем выравнивание текста: для одних ячеек центр, для других justify
                    if cell_num in [0, 1, 3, 4, 5, 7]:
                        cls._style_cell_text(cell)
                    else:
                        cls._style_cell_text(cell, justify=True)

    @classmethod
    def _send_progress_update(cls, send_progress: Optional[Callable[[int], None]], current: int, total: int) -> None:
        """
        Отправляет прогресс выполнения, если передана функция send_progress.

        Args:
            send_progress (Optional[Callable[[int], None]]): Функция для отправки прогресса (от 0 до 100).
            current (int): Текущий номер итерации (например, текущий обработанный элемент).
            total (int): Общее количество элементов.
        """
        progress = int((current / total) * 100)
        logger.info(f"Прогресс создания документа {progress}%")
        if send_progress:
            send_progress(progress)

    @classmethod
    def create_globus(
            cls,
            filename: str = "globus_new.docx",
            send_progress: Optional[Callable[[int], None]] = None,
    ) -> None:
        """
        Создает документ Word с данными о городах и таблицами.

        Args:
            filename (str, optional): Имя файла для сохранения документа. По умолчанию "globus_new.docx".
            send_progress (Optional[Callable[[int], None]], optional): Функция для отправки прогресса создания документа.
        """
        document = cls._init_document()
        cls._create_custom_heading_style(document)
        tables_name = list(TableNames.objects.values_list("id", "table_name"))
        document.add_paragraph("Какой-то текст для начала")

        if not tables_name:
            raise ValueError("Файл пустой. Нет данных для добавления.")

        for num, (table_id, table_name) in enumerate(tables_name):
            paragraph = document.add_paragraph(table_name, style="CustomHeadingStyle")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            document.add_paragraph()

            table_data = list(
                CityData.objects.filter(table_id=table_id).values_list(
                    "table_id",
                    "dock_num",
                    "location",
                    "name_organ",
                    "pseudonim",
                    "letters",
                    "writing",
                    "ip_address",
                    "some_number",
                    "work_time",
                )
            )

            table = cls._create_and_format_table(document, 3 + len(table_data))
            cls._fill_table_headers(table)
            cls._fill_table_data(table, table_data)
            cls._send_progress_update(send_progress, num + 1, len(tables_name))

            document.add_page_break()

        document.save(ProjectSettings.tlg_dir / filename)


if __name__ == "__main__":
    globus_file = ProjectSettings.tlg_dir / 'globus.docx'
    print(globus_file)
    GlobusParser.process_file(file_path=globus_file)
    # print(res)
    GlobusParser.create_globus()
